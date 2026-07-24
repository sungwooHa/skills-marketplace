#!/usr/bin/env python3
"""
범용 마크다운 → 워드(docx) 변환 스크립트
사용법: python3 .claude/scripts/md_to_docx.py <input.md>
출력: 같은 폴더에 .docx 파일 생성

지원: 제목(#/##/###), 테이블, 리스트, 번호 리스트, 인용(>), 인라인 볼드/이탤릭
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# === 설정 ===
FONT_NAME = "맑은 고딕"
FONT_SIZE_BODY = 10
FONT_SIZE_TITLE = 16
FONT_SIZE_H1 = 13
FONT_SIZE_H2 = 11
LINE_SPACING = 1.3

COLOR_TITLE = RGBColor(0x1A, 0x23, 0x7E)
COLOR_H1 = RGBColor(0x15, 0x65, 0xC0)
COLOR_H2 = RGBColor(0x33, 0x33, 0x33)
COLOR_BODY = RGBColor(0x21, 0x21, 0x21)
COLOR_QUOTE = RGBColor(0x55, 0x55, 0x55)
HEADER_BG = "E3F2FD"
TABLE_BORDER_COLOR = "BDBDBD"


def set_font(run, name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text, size=FONT_SIZE_BODY, color=COLOR_BODY,
                       base_bold=False, base_italic=False):
    """인라인 마크다운(**볼드**, *이탤릭*)을 파싱하여 run 단위로 추가"""
    # **bold** 와 *italic* 패턴 처리
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')
    pos = 0
    for m in pattern.finditer(text):
        # 매치 이전 일반 텍스트
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_font(run, size=size, bold=base_bold, color=color, italic=base_italic)
        # 볼드+이탤릭 (***text***)
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            set_font(run, size=size, bold=True, color=color, italic=True)
        # 볼드 (**text**)
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            set_font(run, size=size, bold=True, color=color, italic=base_italic)
        # 이탤릭 (*text*)
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            set_font(run, size=size, bold=base_bold, color=color, italic=True)
        pos = m.end()
    # 나머지 텍스트
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=size, bold=base_bold, color=color, italic=base_italic)
    # 텍스트가 비어있으면 빈 run 하나
    if pos == 0 and not text:
        run = paragraph.add_run("")
        set_font(run, size=size, bold=base_bold, color=color, italic=base_italic)


def add_paragraph(doc, text="", size=FONT_SIZE_BODY, bold=False, color=COLOR_BODY,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = LINE_SPACING
    if text:
        add_formatted_runs(p, text, size=size, color=color,
                           base_bold=bold, base_italic=italic)
    return p


def parse_markdown(md_text):
    """마크다운 텍스트를 섹션 단위로 파싱"""
    lines = md_text.split('\n')
    sections = []
    current = {'type': 'body', 'content': []}

    for line in lines:
        # 제목 (# 회의록: ...)
        if line.startswith('# ') and not line.startswith('## '):
            if current['content']:
                sections.append(current)
            current = {'type': 'title', 'text': line[2:].strip(), 'content': []}
        # H2 (## ...)
        elif line.startswith('## '):
            if current['content'] or current.get('text'):
                sections.append(current)
            current = {'type': 'h1', 'text': line[3:].strip(), 'content': []}
        # H3 (### ...)
        elif line.startswith('### '):
            if current['content'] or current.get('text'):
                sections.append(current)
            current = {'type': 'h2', 'text': line[4:].strip(), 'content': []}
        # 테이블 구분선 (|------|------| 등) — 건너뛰기
        elif re.match(r'^\s*\|[\s\-:]+(\|[\s\-:]+)*\|?\s*$', line):
            pass
        # 테이블 행
        elif '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                if 'table' not in current:
                    current['table'] = []
                current['table'].append(cells)
        # 인용 (> ...) — 연속 인용은 병합
        elif line.startswith('> '):
            quote_text = line[2:].strip()
            if (current['content'] and
                    current['content'][-1]['type'] == 'quote'):
                current['content'][-1]['text'] += '\n' + quote_text
            else:
                current['content'].append({'type': 'quote', 'text': quote_text})
        # 리스트 항목 (- ...)
        elif line.startswith('- ') or line.startswith('  - '):
            indent = 1 if line.startswith('  ') else 0
            text = line.lstrip(' -').strip()
            current['content'].append({'type': 'list', 'text': text, 'indent': indent})
        # 들여쓰기 번호 리스트 (  1. ...)
        elif re.match(r'^\s+(\d+)\.\s', line):
            m = re.match(r'^\s+(\d+)\.\s+(.*)', line)
            current['content'].append({
                'type': 'numlist', 'text': m.group(2).strip(),
                'num': m.group(1), 'indent': 1
            })
        # 번호 리스트 (1. ...)
        elif re.match(r'^(\d+)\.\s', line):
            m = re.match(r'^(\d+)\.\s+(.*)', line)
            current['content'].append({
                'type': 'numlist', 'text': m.group(2).strip(),
                'num': m.group(1), 'indent': 0
            })
        # 일반 텍스트
        elif line.strip():
            current['content'].append({'type': 'text', 'text': line.strip()})

    if current['content'] or current.get('text') or current.get('table'):
        sections.append(current)

    return sections


def build_docx(sections, output_path):
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for sec in sections:
        if sec['type'] == 'title':
            add_paragraph(doc, sec['text'], size=FONT_SIZE_TITLE, bold=True,
                          color=COLOR_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

        elif sec['type'] == 'h1':
            # 구분선
            add_paragraph(doc, "", space_after=2)
            add_paragraph(doc, sec['text'], size=FONT_SIZE_H1, bold=True,
                          color=COLOR_H1, space_after=8)

        elif sec['type'] == 'h2':
            add_paragraph(doc, sec['text'], size=FONT_SIZE_H2, bold=True,
                          color=COLOR_H2, space_after=6)

        # 테이블 처리
        if sec.get('table'):
            rows = sec['table']
            if len(rows) >= 1:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    for j, cell_text in enumerate(row):
                        if j < len(table.columns):
                            cell = table.cell(i, j)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            is_header = (i == 0)
                            add_formatted_runs(p, cell_text, size=FONT_SIZE_BODY,
                                               color=COLOR_BODY, base_bold=is_header)
                            if is_header:
                                set_cell_bg(cell, HEADER_BG)

                add_paragraph(doc, "", space_after=4)

        # 컨텐츠 처리
        for item in sec.get('content', []):
            if item['type'] == 'quote':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.left_indent = Cm(1)
                # 멀티라인 인용 처리 (연속 > 줄 병합됨)
                quote_lines = item['text'].split('\n')
                for li, qline in enumerate(quote_lines):
                    if li > 0:
                        run = p.add_run('\n')
                        set_font(run, size=FONT_SIZE_BODY, color=COLOR_QUOTE)
                    add_formatted_runs(p, qline, size=FONT_SIZE_BODY,
                                       color=COLOR_QUOTE, base_italic=True)

            elif item['type'] == 'list':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = LINE_SPACING
                indent = Cm(0.5) if item['indent'] == 0 else Cm(1.2)
                p.paragraph_format.left_indent = indent
                bullet = "\u2022 " if item['indent'] == 0 else "\u25E6 "
                run = p.add_run(bullet)
                set_font(run, size=FONT_SIZE_BODY, color=COLOR_BODY)
                add_formatted_runs(p, item['text'], size=FONT_SIZE_BODY, color=COLOR_BODY)

            elif item['type'] == 'numlist':
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = LINE_SPACING
                indent = Cm(0.5) if item.get('indent', 0) == 0 else Cm(1.2)
                p.paragraph_format.left_indent = indent
                num = item.get('num', '1')
                run = p.add_run(f"{num}. ")
                set_font(run, size=FONT_SIZE_BODY, color=COLOR_BODY, bold=False)
                add_formatted_runs(p, item['text'], size=FONT_SIZE_BODY, color=COLOR_BODY)

            elif item['type'] == 'text':
                add_paragraph(doc, item['text'], space_after=4)

    doc.save(output_path)
    print(f"생성 완료: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("사용법: python3 md_to_docx.py <input.md>")
        sys.exit(1)

    md_path = sys.argv[1]
    if not os.path.exists(md_path):
        print(f"파일을 찾을 수 없습니다: {md_path}")
        sys.exit(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    output_path = os.path.splitext(md_path)[0] + '.docx'
    sections = parse_markdown(md_text)
    build_docx(sections, output_path)


if __name__ == '__main__':
    main()
